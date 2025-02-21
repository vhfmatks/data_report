"""
결과 보고서 생성 모듈
"""

from datetime import datetime
from typing import Dict, Any, List, Tuple
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np

def save_plot_to_image():
    """현재 matplotlib 차트를 이미지로 저장"""
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
    img_stream.seek(0)
    plt.close()
    return img_stream

def generate_report(
    df,
    schema: Dict,
    analysis_results: Dict,
    insights: str,
    visualizations: Dict[str, List[Tuple[str, Any]]] = None
) -> bytes:
    """결과 보고서 생성"""
    try:
        report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Word 문서 생성
        doc = Document()
        
        # 제목 스타일 설정
        title = doc.add_heading('데이터 분석 결과 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성일시
        doc.add_paragraph(f'생성일시: {report_date}')
        doc.add_paragraph()
        
        # 1. 개요
        doc.add_heading('1. 개요', 1)
        overview = doc.add_paragraph()
        overview.add_run(f'• 분석 데이터셋: {df.shape[0]:,}행 × {df.shape[1]:,}열\n')
        overview.add_run('• 데이터 유형:\n')
        overview.add_run(f'  - 수치형 변수: {len([col for col, info in schema.items() if info["data_type"] == "numeric"])}개\n')
        overview.add_run(f'  - 범주형 변수: {len([col for col, info in schema.items() if info["data_type"] == "categorical"])}개\n')
        overview.add_run(f'  - 시계열 변수: {len([col for col, info in schema.items() if info["data_type"] == "datetime"])}개')
        
        # 2. 주요 발견사항
        doc.add_heading('2. 주요 발견사항', 1)
        doc.add_paragraph(insights)
        
        # 3. 변수별 상세 분석
        doc.add_heading('3. 변수별 상세 분석', 1)
        
        for col, info in schema.items():
            if col not in analysis_results:
                continue
                
            display_name = info.get('display_name', col)
            doc.add_heading(f'{display_name}({col})', 2)
            
            # 변수 설명
            description = doc.add_paragraph()
            description.add_run('설명: ').bold = True
            description.add_run(info.get('description', '설명 없음'))
            
            data_type = doc.add_paragraph()
            data_type.add_run('데이터 타입: ').bold = True
            data_type.add_run(info['data_type'])
            
            # 데이터 타입별 상세 정보 추가
            try:
                if info['data_type'] == 'numeric':
                    stats = analysis_results[col].get("기본통계", {})
                    outliers = analysis_results[col].get("이상치", {})
                    
                    doc.add_heading('기본 통계', 3)
                    stats_table = doc.add_table(rows=1, cols=2)
                    stats_table.style = 'Table Grid'
                    
                    # 헤더 추가
                    header_cells = stats_table.rows[0].cells
                    header_cells[0].text = '지표'
                    header_cells[1].text = '값'
                    
                    # 통계값 추가
                    stats_data = [
                        ('평균', f"{stats.get('평균', 0):.2f}"),
                        ('중앙값', f"{stats.get('중앙값', 0):.2f}"),
                        ('표준편차', f"{stats.get('표준편차', 0):.2f}"),
                        ('최소값', f"{stats.get('최소값', 0):.2f}"),
                        ('최대값', f"{stats.get('최대값', 0):.2f}")
                    ]
                    
                    for name, value in stats_data:
                        row_cells = stats_table.add_row().cells
                        row_cells[0].text = name
                        row_cells[1].text = value
                    
                    # 분포 시각화 추가
                    plt.figure(figsize=(10, 6))
                    sns.histplot(data=df[col], bins=30, kde=True)
                    plt.title(f"{display_name} 분포")
                    plt.xlabel(display_name)
                    plt.ylabel("빈도")
                    doc.add_picture(save_plot_to_image(), width=Inches(6))
                    
                    # 박스플롯 추가
                    plt.figure(figsize=(8, 4))
                    sns.boxplot(data=df[col])
                    plt.title(f"{display_name} 박스플롯")
                    doc.add_picture(save_plot_to_image(), width=Inches(6))
                    
                    doc.add_paragraph()
                    doc.add_heading('이상치 정보', 3)
                    outlier_info = doc.add_paragraph()
                    outlier_info.add_run(f'• 이상치 개수: {outliers.get("개수", 0)}\n')
                    outlier_info.add_run(f'• 이상치 비율: {outliers.get("비율", 0):.2f}%')
                    
                elif info['data_type'] == 'categorical':
                    value_counts = analysis_results[col].get("고유값", {})
                    mode = analysis_results[col].get("최빈값", {})
                    
                    doc.add_heading('범주 분포', 3)
                    cat_table = doc.add_table(rows=1, cols=3)
                    cat_table.style = 'Table Grid'
                    
                    # 헤더 추가
                    header_cells = cat_table.rows[0].cells
                    header_cells[0].text = '범주'
                    header_cells[1].text = '빈도'
                    header_cells[2].text = '비율'
                    
                    # 상위 5개 범주 추가
                    distribution = value_counts.get('분포', {})
                    for category, stats in list(distribution.items())[:5]:
                        row_cells = cat_table.add_row().cells
                        row_cells[0].text = str(category)
                        row_cells[1].text = f"{stats.get('빈도', 0):,}"
                        row_cells[2].text = f"{stats.get('비율', 0):.1f}%"
                    
                    # 막대 그래프 추가
                    plt.figure(figsize=(10, 6))
                    categories = list(distribution.keys())[:5]
                    frequencies = [distribution[cat]['빈도'] for cat in categories]
                    plt.bar(categories, frequencies)
                    plt.title(f"{display_name} 상위 5개 범주 분포")
                    plt.xticks(rotation=45)
                    plt.xlabel("범주")
                    plt.ylabel("빈도")
                    doc.add_picture(save_plot_to_image(), width=Inches(6))
                    
                    # 원형 차트 추가
                    plt.figure(figsize=(8, 8))
                    plt.pie(frequencies, labels=categories, autopct='%1.1f%%')
                    plt.title(f"{display_name} 범주 비율")
                    doc.add_picture(save_plot_to_image(), width=Inches(6))
                    
                elif info['data_type'] == 'datetime':
                    period = analysis_results[col].get("기간", {})
                    
                    doc.add_heading('시간 정보', 3)
                    time_info = doc.add_paragraph()
                    time_info.add_run(f'• 시작: {period.get("시작", "N/A")}\n')
                    time_info.add_run(f'• 종료: {period.get("종료", "N/A")}\n')
                    time_info.add_run(f'• 총 일수: {period.get("기간", {}).get("일수", 0)}일')
                    
                    # 시계열 트렌드 추가
                    if "분포" in analysis_results[col]:
                        # 월별 분포 차트
                        monthly_dist = analysis_results[col]["분포"].get("월별", {})
                        if monthly_dist:
                            plt.figure(figsize=(12, 6))
                            months = list(monthly_dist.keys())
                            counts = list(monthly_dist.values())
                            plt.plot(months, counts, marker='o')
                            plt.title(f"{display_name} 월별 분포")
                            plt.xticks(rotation=45)
                            plt.xlabel("월")
                            plt.ylabel("빈도")
                            plt.grid(True)
                            doc.add_picture(save_plot_to_image(), width=Inches(6))
                    
            except Exception as e:
                doc.add_paragraph(f'⚠️ 상세 정보 생성 중 오류 발생: {str(e)}')
        
        # 4. 상관관계 분석
        if "상관관계" in analysis_results and analysis_results["상관관계"]:
            doc.add_heading('4. 상관관계 분석', 1)
            
            # 상관관계 표 추가
            corr_table = doc.add_table(rows=1, cols=3)
            corr_table.style = 'Table Grid'
            
            # 헤더 추가
            header_cells = corr_table.rows[0].cells
            header_cells[0].text = '변수 1'
            header_cells[1].text = '변수 2'
            header_cells[2].text = '상관계수'
            
            # 상관관계 정보 추가
            for corr in analysis_results["상관관계"]:
                row_cells = corr_table.add_row().cells
                row_cells[0].text = corr['변수1']
                row_cells[1].text = corr['변수2']
                row_cells[2].text = f"{corr['상관계수']:.3f}"
            
            # 상관관계 히트맵 추가
            numeric_cols = [col for col, info in schema.items() 
                          if info["data_type"] == "numeric" and col in df.columns]
            if len(numeric_cols) > 1:
                plt.figure(figsize=(10, 8))
                correlation_matrix = df[numeric_cols].corr()
                sns.heatmap(correlation_matrix, 
                           annot=True, 
                           cmap='RdYlBu_r',
                           center=0,
                           vmin=-1,
                           vmax=1,
                           fmt='.2f')
                plt.title("변수 간 상관관계 히트맵")
                doc.add_picture(save_plot_to_image(), width=Inches(6))
        
        # Word 문서를 바이트로 변환
        doc_binary = io.BytesIO()
        doc.save(doc_binary)
        doc_binary.seek(0)
        
        return doc_binary.getvalue()
        
    except Exception as e:
        # 오류 발생 시 빈 바이트 반환
        return b"" 